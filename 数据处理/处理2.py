from docx import Document
import glob
import os
import csv
import docx

files = glob.glob(os.path.join("*.docx"))
# files = ["SCFF.INFO(M).MT.047 金箔入酒事件 （2015年3月）.docx"]
log_path = '数据提取.csv'
file = open(log_path, 'a+', encoding='utf-8_sig', newline='')
csv_writer = csv.writer(file)
csv_writer.writerow([f'标题', '文本', '文件名'])

def get_paragraphs(docx_path):
    title = ""
    text = ""
    # 打开word文档
    document = Document(docx_path)

    # 获取所有段落
    all_paragraphs = document.paragraphs
    paragraph_texts = []

    flag = 0
    # 循环读取列表
    for p in document.paragraphs:
        a = p.style.style_id
        if flag == 1:
            if p.style.name == 'Normal':
                if len(p.runs) and p.runs[0].font.size is not None:
                    if p.runs[0].font.size.pt > 12.1:
                        flag = 0
                        text = text.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ').replace(u'&nbsp;', ' ')
                        row = [title, text, docx_path[:-5]]
                        csv_writer.writerow(row)
                        text = ""
                    else:
                        text += p.text
                else:
                    text += p.text
            else:
                flag = 0
                text = text.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ').replace(u'&nbsp;', ' ')
                row = [title, text, docx_path[:-5]]
                csv_writer.writerow(row)
                text = ""
        if p.style.name == 'Heading 3':
            title = p.text
            flag = 1
    return paragraph_texts


def spaceReplace(i):
    i = i.replace('  ', ' ')
    i = spaceReplace(i) if '  ' in i else i
    return i




for filename in files:
    docx_path = filename
    paragraph_texts = get_paragraphs(docx_path)
file.close()
